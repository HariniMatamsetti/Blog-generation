import torch
import streamlit as st
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
from diffusers import StableDiffusionPipeline
from docx import Document
from PIL import Image
import io

# Define model names
TEXT_MODEL = "meta-llama/Llama-2-7b-chat-hf"  # Use chat-tuned version
IMAGE_MODEL = "runwayml/stable-diffusion-v1-5"

# Load text generation model with optimized settings
@st.cache_resource
@st.cache_resource
def load_text_model():
    tokenizer = AutoTokenizer.from_pretrained(TEXT_MODEL)
    model = AutoModelForCausalLM.from_pretrained(
        TEXT_MODEL,
        torch_dtype=torch.float16,  # Use FP16 instead of bitsandbytes 4-bit
        device_map="auto"  # Auto-assign model to GPU
    )
    return model, tokenizer


# Load image generation model with improved quality
@st.cache_resource
def load_image_model():
    pipe = StableDiffusionPipeline.from_pretrained(
        IMAGE_MODEL,
        torch_dtype=torch.float16,
        device_map="balanced"
    )
    return pipe

# Load models
model, tokenizer = load_text_model()
pipe = load_image_model()

# Streamlit UI
st.title("AI Text & Image Generator")

# Text & Image Input Fields
text_prompt = st.text_input("Enter text prompt for Llama 2:")
image_prompt = st.text_input("Enter image prompt for Stable Diffusion:")

# Generate button
if st.button("Generate"):
    # Generate text
    with st.spinner("Generating text..."):
        chat_prompt = f"[INST] {text_prompt}. Please respond only in clear, fluent English. [/INST]"
        inputs = tokenizer(chat_prompt, return_tensors="pt").to(model.device)
        output = model.generate(
            **inputs,
            max_new_tokens=300,
            temperature=0.6,
            top_k=40,
            top_p=0.85,
            repetition_penalty=1.3
        )
        generated_text = tokenizer.decode(output[0], skip_special_tokens=True)
        st.subheader("Generated Text:")
        st.write(generated_text)

    # Generate image
    with st.spinner("Generating image..."):
        image = pipe(image_prompt, height=512, width=512).images[0]  # Increased resolution
        st.subheader("Generated Image:")
        st.image(image, caption="Generated Image", use_column_width=True)

    # Add Download Option
    if st.button("Download as Word Document"):
        doc = Document()
        doc.add_heading("AI Generated Content", level=1)

        # Add text
        doc.add_heading("Generated Text:", level=2)
        doc.add_paragraph(generated_text)

        # Add image
        image_stream = io.BytesIO()
        image.save(image_stream, format="PNG")
        doc.add_picture(image_stream, width=4000000)  # Scales image

        # Save file
        doc_path = "generated_content.docx"
        doc.save(doc_path)
        with open(doc_path, "rb") as file:
            st.download_button("Download Word Document", file, file_name="generated_content.docx")

