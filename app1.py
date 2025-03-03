import torch
import streamlit as st
from transformers import AutoModelForCausalLM, AutoTokenizer, BitsAndBytesConfig
from diffusers import StableDiffusionPipeline
from PIL import Image
from io import BytesIO
from docx import Document

# Define model names
TEXT_MODEL = "meta-llama/Llama-2-7b-hf"
IMAGE_MODEL = "runwayml/stable-diffusion-v1-5"

# Load text generation model with 4-bit quantization
@st.cache_resource
def load_text_model():
    bnb_config = BitsAndBytesConfig(
        load_in_4bit=True,
        bnb_4bit_compute_dtype=torch.float16,
        bnb_4bit_use_double_quant=True
    )
    tokenizer = AutoTokenizer.from_pretrained(TEXT_MODEL)
    model = AutoModelForCausalLM.from_pretrained(
        TEXT_MODEL,
        quantization_config=bnb_config,
        device_map="balanced"
    )
    return model, tokenizer

# Load image generation model
@st.cache_resource
def load_image_model():
    pipe = StableDiffusionPipeline.from_pretrained(
        IMAGE_MODEL,
        torch_dtype=torch.float16,
        device_map="balanced"
    )
    return pipe

# Load models
model, tokenizer = load_text_model()
pipe = load_image_model()

# Streamlit UI
st.title("📝🎨 AI Text & Image Generator")

# Session state to store outputs
if "generated_text" not in st.session_state:
    st.session_state.generated_text = None
if "generated_images" not in st.session_state:
    st.session_state.generated_images = []

# Text Generation Input
prompt = st.text_input("Enter text prompt for Llama 2:")
if st.button("Generate Text"):
    with st.spinner("Generating text..."):
        inputs = tokenizer(prompt, return_tensors="pt").to(model.device)
        output = model.generate(**inputs, max_new_tokens=150)  # Increased output length
        generated_text = tokenizer.decode(output[0], skip_special_tokens=True)

        # Store generated text
        st.session_state.generated_text = generated_text
        st.write("### Generated Text:")
        st.write(generated_text)

# Image Generation Input
img_prompt = st.text_input("Enter image prompt for Stable Diffusion:")
num_images = st.slider("Select number of images to generate", 1, 5, 1)  # User selects 1-5 images

if st.button("Generate Image(s)"):
    with st.spinner(f"Generating {num_images} image(s)..."):
        images = [pipe(img_prompt, height=512, width=512).images[0] for _ in range(num_images)]  # Generate multiple images
        st.session_state.generated_images = images  # Store images in session state

        # Display all generated images
        for idx, image in enumerate(images):
            st.image(image, caption=f"Generated Image {idx+1}", use_column_width=True)

# Download Button
if st.session_state.generated_text or st.session_state.generated_images:
    st.write("### Download Generated Content")
    if st.button("Download Word Document"):
        doc = Document()

        # Add text to Word document
        if st.session_state.generated_text:
            doc.add_heading("Generated Text", level=1)
            doc.add_paragraph(st.session_state.generated_text)

        # Add images to Word document
        if st.session_state.generated_images:
            doc.add_heading("Generated Images", level=1)
            for idx, image in enumerate(st.session_state.generated_images):
                image_stream = BytesIO()
                image.save(image_stream, format="PNG")
                doc.add_paragraph(f"Image {idx+1}:")
                doc.add_picture(image_stream)

        # Save document
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Provide download button
        st.download_button(label="📄 Download Word File", data=doc_io, file_name="AI_Generated.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
